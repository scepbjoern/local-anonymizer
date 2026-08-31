"""Document extraction utilities for various file formats."""

from __future__ import annotations

import atexit
import concurrent.futures
import csv
import io
import json
import logging
import os
import re
import sys
import threading
import time
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Sequence, Set, Tuple, Union

_pdf_env_lock = threading.Lock()

CONFIG_DIR = Path.home() / ".local-anonymizer"
TEMP_UPLOADS_DIR = CONFIG_DIR / "temp_uploads"
TEMP_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)


def is_pdf_worker() -> bool:
    """Returns True if the current process is a spawned PDF extraction worker sub-process."""
    return os.environ.get("LOCAL_ANONYMIZER_PDF_WORKER") == "1"


def cleanup_extraction_temp_files(max_age_seconds: int = 1800):
    """Clean up any stale temporary PDF extraction files older than max_age_seconds (default 30 min).
    Never deletes recently active or newly created files from other running app instances or tabs."""
    if is_pdf_worker():
        return
    try:
        if TEMP_UPLOADS_DIR.exists():
            cutoff = time.time() - max_age_seconds
            for f in TEMP_UPLOADS_DIR.glob("*.pdf"):
                try:
                    if f.stat().st_mtime < cutoff:
                        f.unlink(missing_ok=True)
                except Exception:
                    pass
    except Exception:
        pass


if not is_pdf_worker():
    cleanup_extraction_temp_files()
    atexit.register(cleanup_extraction_temp_files)


class UnsupportedFileFormatError(ValueError):
    """Raised when an unsupported file format is provided to the extractor."""
    pass


def strip_html_markup(text: str) -> str:
    """Strip formatting HTML tags (<mark>, <u>, <ins>, <span>, <font>, <del>, <strike>) while preserving the inner text."""
    if not text:
        return ""
    return re.sub(r"</?(?:mark|ins|u|span|font|strike|del)(?:\s+[^>]*)?>", "", text, flags=re.IGNORECASE).strip()


def extract_text_from_txt_bytes(raw_bytes: bytes) -> str:
    """
    Extract text from plain text or markdown bytes with robust encoding detection.
    Tries UTF-8 (with BOM), UTF-8, Windows CP1252, ISO-8859-15, and Latin-1 in order.
    """
    encodings = ["utf-8-sig", "utf-8", "cp1252", "iso-8859-15", "latin-1"]
    for enc in encodings:
        try:
            return raw_bytes.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    # Ultimate fallback with character replacement
    return raw_bytes.decode("utf-8", errors="replace")


def safe_read_bytes(file_path: Union[str, Path]) -> bytes:
    """
    Safely read bytes from a file path. On Windows, uses Win32 API with full sharing
    (FILE_SHARE_READ | FILE_SHARE_WRITE | FILE_SHARE_DELETE) so that files currently open
    in Microsoft Word, Excel, OneDrive, or other editors can be read without PermissionError.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Datei nicht gefunden: {path}")

    if sys.platform == "win32":
        try:
            import ctypes
            from ctypes import wintypes

            kernel32 = ctypes.windll.kernel32
            GENERIC_READ = 0x80000000
            FILE_SHARE_READ = 1
            FILE_SHARE_WRITE = 2
            FILE_SHARE_DELETE = 4
            OPEN_EXISTING = 3
            FILE_ATTRIBUTE_NORMAL = 0x80
            INVALID_HANDLE_VALUE = -1

            handle = kernel32.CreateFileW(
                str(path.resolve()),
                GENERIC_READ,
                FILE_SHARE_READ | FILE_SHARE_WRITE | FILE_SHARE_DELETE,
                None,
                OPEN_EXISTING,
                FILE_ATTRIBUTE_NORMAL,
                None,
            )
            if handle != INVALID_HANDLE_VALUE:
                try:
                    size = kernel32.GetFileSize(handle, None)
                    if size > 0:
                        buf = ctypes.create_string_buffer(size)
                        bytes_read = wintypes.DWORD()
                        if kernel32.ReadFile(handle, buf, size, ctypes.byref(bytes_read), None):
                            return buf.raw[:bytes_read.value]
                finally:
                    kernel32.CloseHandle(handle)
        except Exception as e:
            logging.debug(f"Win32 safe_read_bytes fallback: {e}")

    # Standard fallback (macOS, Linux, or if Win32 wasn't used)
    return path.read_bytes()


def extract_text_from_txt(path: Path) -> str:
    """Extract text from plain text or markdown files."""
    return extract_text_from_txt_bytes(safe_read_bytes(path))


def extract_text_from_json_bytes(data: bytes) -> str:
    """Extract JSON bytes as cleanly formatted JSON text."""
    raw_str = extract_text_from_txt_bytes(data)
    try:
        parsed = json.loads(raw_str)
        return json.dumps(parsed, indent=2, ensure_ascii=False)
    except Exception:
        return raw_str


def clean_markdown_table_cell(cell: str) -> str:
    """
    Format and clean text inside Markdown table cells:
    - Reconnects broken hyphenations ('Dokumenten-<br>scanning' -> 'Dokumentenscanning', 'Sach-<br>bearbeiter:in' -> 'Sachbearbeiter:in').
    - Preserves hyphenated conjunctions ('Tarif-<br>und' -> 'Tarif- und').
    - Fixes broken bullet lists ('-<br>Veranlasser:in' -> '- Veranlasser:in').
    - Replaces soft intra-sentence linebreaks with a clean space while preserving paragraph breaks (<br><br>).
    - Removes trailing orphan dashes or empty bullet points.
    """
    c = cell.strip()
    if not c or c == "-":
        return ""

    # 1. Protect hyphenated conjunctions: 'Tarif- <br> und' -> 'Tarif- und'
    c = re.sub(r"([A-Za-zÄÖÜäöüß]+-)\s*(?:<br\s*/?>|\n)\s*(und|oder|bzw|sowie)\b", r"\1 \2", c, flags=re.IGNORECASE)

    # 2. Fix broken hyphenated word wraps: 'Dokumenten-<br>scanning' -> 'Dokumentenscanning'
    c = re.sub(r"([A-Za-zÄÖÜäöüß]+)-\s*(?:<br\s*/?>|\n)\s*([a-zäöüß]+)", r"\1\2", c)
    c = re.sub(r"([A-Za-zÄÖÜäöüß]+)-\s*(?:<br\s*/?>|\n)\s*([A-ZÄÖÜ][a-zäöüß]+)", r"\1-\2", c)

    # 3. Normalize paragraph double linebreaks
    c = re.sub(r"(?:<br\s*/?>\s*){2,}", " __PARAGRAPH_BREAK__ ", c)

    # 4. Fix broken bullet points: '-<br>Text' -> '- Text'
    c = re.sub(r"(?:^|<br\s*/?>|\n)\s*[-•–—]\s*(?:<br\s*/?>|\n)\s*", " __BULLET__ ", c)
    c = re.sub(r"(?:^|<br\s*/?>|\n)\s*[-•–—]\s+", " __BULLET__ ", c)

    # 5. Replace single soft linebreaks within flow text with a space
    c = re.sub(r"<br\s*/?>|\n", " ", c)

    # 6. Reconstruct bullets and paragraphs
    c = c.replace("__BULLET__", "<br>- ")
    c = c.replace("__PARAGRAPH_BREAK__", "<br><br>")

    # 7. Strip trailing empty bullets / dashes (e.g. '<br>- <br>- ')
    c = re.sub(r"(?:<br\s*/?>\s*[-•–—\s]*)+$", "", c, flags=re.IGNORECASE)
    c = re.sub(r"[\s\-•–—]+$", "", c)
    c = re.sub(r"^(?:<br\s*/?>|\s)*", "", c)
    c = re.sub(r"[ \t]+", " ", c)

    return c.strip()


def extract_text_from_csv_bytes(data: bytes) -> str:
    """Extract CSV bytes as a structured Markdown table, supporting multi-line quoted cells."""
    text_content = extract_text_from_txt_bytes(data)
    if not text_content.strip():
        return ""

    # Detect delimiter: try comma, semicolon, tab
    sample_lines = [l for l in text_content.splitlines() if l.strip()][:5]
    sample = "\n".join(sample_lines)
    delimiter = ","
    try:
        sniffer = csv.Sniffer()
        dialect = sniffer.sniff(sample, delimiters=";,|\t")
        delimiter = dialect.delimiter
    except Exception:
        if ";" in sample and "," not in sample:
            delimiter = ";"
        elif "\t" in sample:
            delimiter = "\t"

    reader = csv.reader(io.StringIO(text_content), delimiter=delimiter)
    rows = [r for r in reader if any(c.strip() for c in r)]
    if not rows:
        return text_content

    max_cols = max(len(r) for r in rows)
    if max_cols == 0:
        return text_content

    def format_csv_cell(cell: str) -> str:
        c = cell.strip()
        c = c.replace("\r\n", "\n").replace("\r", "\n")
        lines = [line.strip() for line in c.split("\n") if line.strip()]
        c = "<br>".join(lines)
        c = c.replace("|", r"\|")
        return c

    # Normalize row lengths and format cells
    norm_rows = []
    for r in rows:
        padded = [format_csv_cell(c) for c in r]
        while len(padded) < max_cols:
            padded.append("")
        norm_rows.append(padded)

    header = norm_rows[0]
    md_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join([":---"] * max_cols) + " |",
    ]
    for r in norm_rows[1:]:
        md_lines.append("| " + " | ".join(r) + " |")

    return "\n".join(md_lines)


def _detect_docx_list_type(para, doc) -> Optional[str]:
    """
    Detect whether a paragraph is a bullet list, numbered list, or normal text.
    Returns 'bullet', 'number', or None.
    """
    style_name = para.style.name.lower() if para.style and para.style.name else ""
    if any(k in style_name for k in ["bullet", "listenpunkt", "aufzählung", "aufzaehlung", "list bullet"]):
        return "bullet"
    if any(k in style_name for k in ["list number", "nummerierung", "list 2", "list 3"]):
        return "number"

    # Inspect XML w:numPr
    num_pr = para._element.xpath("./w:pPr/w:numPr")
    if not num_pr:
        return None

    try:
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        num_id_nodes = para._element.xpath("./w:pPr/w:numPr/w:numId/@w:val")
        ilvl_nodes = para._element.xpath("./w:pPr/w:numPr/w:ilvl/@w:val")
        num_id = num_id_nodes[0] if num_id_nodes else None
        ilvl = ilvl_nodes[0] if ilvl_nodes else "0"

        if num_id and hasattr(doc, "part") and hasattr(doc.part, "numbering_part") and doc.part.numbering_part is not None:
            np = doc.part.numbering_part._element
            num_nodes = np.xpath(f'./w:num[@w:numId="{num_id}"]/w:abstractNumId/@w:val', namespaces=ns)
            if num_nodes:
                ab_id = num_nodes[0]
                fmt_nodes = np.xpath(f'./w:abstractNum[@w:abstractNumId="{ab_id}"]/w:lvl[@w:ilvl="{ilvl}"]/w:numFmt/@w:val', namespaces=ns)
                if fmt_nodes:
                    fmt = fmt_nodes[0].lower()
                    if fmt == "bullet":
                        return "bullet"
                    elif fmt in ["decimal", "decimalzero", "lowerletter", "upperletter", "lowerroman", "upperroman", "ordinal", "chinesecounting", "cardinaltext", "ordinaltext", "hex"]:
                        return "number"
    except Exception:
        pass

    # Default for list paragraphs where style indicates number
    if "number" in style_name or "nummer" in style_name:
        return "number"
    return "bullet"


_INLINE_MD_PATTERN = re.compile(
    r"(\*\*\*[^*]+\*\*\*|___[^_]+___|"
    r"\*\*[^*]+\*\*|__[^_]+__|"
    r"\*[^*]+\*|(?<!\w)_[^_]+_(?!\w)|"
    r"`[^`]+`)"
)


def _add_formatted_markdown_runs(para, text: str) -> None:
    """
    Parse inline markdown tokens (bold, italic, bold-italic, code) and add corresponding docx runs
    without literal Markdown syntax markers (**...**, *...*, `...`).
    """
    parts = _INLINE_MD_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if (part.startswith("***") and part.endswith("***") and len(part) >= 6) or \
           (part.startswith("___") and part.endswith("___") and len(part) >= 6):
            content = part[3:-3]
            r = para.add_run(content)
            r.bold = True
            r.italic = True
        elif (part.startswith("**") and part.endswith("**") and len(part) >= 4) or \
             (part.startswith("__") and part.endswith("__") and len(part) >= 4):
            content = part[2:-2]
            r = para.add_run(content)
            r.bold = True
        elif (part.startswith("*") and part.endswith("*") and len(part) >= 2) or \
             (part.startswith("_") and part.endswith("_") and len(part) >= 2):
            content = part[1:-1]
            r = para.add_run(content)
            r.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            content = part[1:-1]
            r = para.add_run(content)
            r.font.name = "Consolas"
        else:
            para.add_run(part)


def extract_text_from_docx_bytes(raw_bytes: bytes, include_headers_footers: bool = False) -> str:
    """
    Extract text from Microsoft Word .docx bytes with:
    - Level 1: XML Outline Level (w:outlineLvl 0..5) for custom corporate heading templates
    - Level 2: Multilingual style names (Heading 1..4, Überschrift 1..4, Title, Subtitle)
    - Level 3: Bullet & numbered lists via XML numbering resolution (w:numFmt) and list styles
    - Level 4: Strict false-positive protection for bold text (only < 60 chars, no ending period, >= 16pt)
    - Level 5: Markdown tables
    - Optional: Header & footer extraction (include_headers_footers=True)
    """
    import docx
    doc = docx.Document(io.BytesIO(raw_bytes))
    full_text = []

    # Optional headers extraction
    if include_headers_footers and doc.sections:
        for sec in doc.sections:
            for p in sec.header.paragraphs:
                if p.text and p.text.strip():
                    full_text.append(p.text.strip())

    for para in doc.paragraphs:
        if not para.text or not para.text.strip():
            continue
        text = para.text.strip()
        style_name = para.style.name.lower() if para.style and para.style.name else ""

        # Stufe 1: XML Gliederungsebene w:outlineLvl
        outline_val = None
        try:
            outline_nodes = para._element.xpath("./w:pPr/w:outlineLvl/@w:val")
            if outline_nodes:
                outline_val = int(outline_nodes[0])
        except Exception:
            pass

        # Stufe 2: Präzise Listen-Erkennung (Bullet vs. Nummerierung)
        list_type = _detect_docx_list_type(para, doc)

        if outline_val is not None and 0 <= outline_val <= 5:
            prefix = "#" * (outline_val + 1)
            text = f"{prefix} {text}"
        elif "heading 1" in style_name or "überschrift 1" in style_name or style_name in ["title", "titel"]:
            text = f"# {text}"
        elif "heading 2" in style_name or "überschrift 2" in style_name or style_name in ["subtitle", "untertitel"]:
            text = f"## {text}"
        elif "heading 3" in style_name or "überschrift 3" in style_name:
            text = f"### {text}"
        elif "heading 4" in style_name or "überschrift 4" in style_name:
            text = f"#### {text}"
        elif list_type == "bullet":
            text = f"- {text}"
        elif list_type == "number":
            text = f"1. {text}"
        elif "quote" in style_name or "zitat" in style_name:
            text = f"> {text}"
        else:
            # Stufe 4: Strict bold-as-heading detection
            # Only treat as heading if:
            # a) All runs are bold
            # b) Total length is short (< 60 chars)
            # c) Does NOT end with punctuation (., !, ?)
            # d) Font size is larger than standard body text (>= 16pt / 32 half-points)
            is_all_bold = (
                len(para.runs) > 0
                and all(r.bold for r in para.runs if r.text and r.text.strip())
            )
            has_no_sentence_end = not text.endswith((".", "!", "?"))
            is_short = len(text) < 60

            is_large_font = False
            for r in para.runs:
                if r.font and r.font.size and r.font.size.pt >= 16:
                    is_large_font = True
                    break

            if is_all_bold and is_short and has_no_sentence_end and is_large_font:
                text = f"## {text}"
            else:
                # Retain inline bold, italic, code for body text
                formatted_runs = []
                for r in para.runs:
                    r_text = r.text
                    if not r_text:
                        continue
                    if r.bold and r.italic:
                        formatted_runs.append(f"***{r_text.strip()}***" if r_text.strip() else r_text)
                    elif r.bold:
                        formatted_runs.append(f"**{r_text.strip()}**" if r_text.strip() else r_text)
                    elif r.italic:
                        formatted_runs.append(f"*{r_text.strip()}*" if r_text.strip() else r_text)
                    else:
                        formatted_runs.append(r_text)
                if formatted_runs:
                    text = "".join(formatted_runs).strip()

        full_text.append(text)

    # Stufe 5: Tabellen als Markdown
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cell_texts = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cell_texts):
                rows_text.append(cell_texts)
        if rows_text:
            header = rows_text[0]
            table_md = ["| " + " | ".join(header) + " |", "| " + " | ".join([":---"] * len(header)) + " |"]
            for r in rows_text[1:]:
                while len(r) < len(header):
                    r.append("")
                table_md.append("| " + " | ".join(r[:len(header)]) + " |")
            full_text.append("\n".join(table_md))

    # Optional footers extraction
    if include_headers_footers and doc.sections:
        for sec in doc.sections:
            for p in sec.footer.paragraphs:
                if p.text and p.text.strip():
                    full_text.append(p.text.strip())

    res_text = "\n\n".join(full_text)
    return strip_html_markup(res_text)


def extract_text_from_docx(path: Path, include_headers_footers: bool = False) -> str:
    """Extract text from Microsoft Word .docx files."""
    return extract_text_from_docx_bytes(safe_read_bytes(path), include_headers_footers=include_headers_footers)


def create_docx_from_markdown(md_text: str):
    """
    Convert Markdown text into a native Word .docx document using real Word paragraph styles
    (Heading 1, Heading 2, Heading 3, Heading 4, List Bullet, List Number, Normal)
    and parse inline formatting (**bold**, *italic*, `code`) into true Word run styles.
    Does NOT output literal '#' or Markdown markers into the document.
    """
    import docx
    doc = docx.Document()
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        trimmed = line.strip()

        if not trimmed:
            i += 1
            continue

        # Headings
        if trimmed.startswith("#### "):
            p = doc.add_heading(level=4)
            _add_formatted_markdown_runs(p, trimmed[5:].strip())
        elif trimmed.startswith("### "):
            p = doc.add_heading(level=3)
            _add_formatted_markdown_runs(p, trimmed[4:].strip())
        elif trimmed.startswith("## "):
            p = doc.add_heading(level=2)
            _add_formatted_markdown_runs(p, trimmed[3:].strip())
        elif trimmed.startswith("# "):
            p = doc.add_heading(level=1)
            _add_formatted_markdown_runs(p, trimmed[2:].strip())
        # Lists
        elif trimmed.startswith("- ") or trimmed.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_markdown_runs(p, trimmed[2:].strip())
        elif re.match(r"^\d+\.\s+", trimmed):
            text_part = re.sub(r"^\d+\.\s+", "", trimmed)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_markdown_runs(p, text_part.strip())
        # Blockquotes
        elif trimmed.startswith("> "):
            p = doc.add_paragraph(style="Quote" if "Quote" in doc.styles else "Normal")
            _add_formatted_markdown_runs(p, trimmed[2:].strip())
        # Tables (Markdown | col1 | col2 |)
        elif trimmed.startswith("|") and trimmed.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1

            rows_data = []
            for t_line in table_lines:
                cells = [c.strip() for c in t_line.strip("|").split("|")]
                if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                    continue
                rows_data.append(cells)

            if rows_data:
                num_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=num_cols)
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass
                for r_idx, r_data in enumerate(rows_data):
                    for c_idx, cell_val in enumerate(r_data):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _add_formatted_markdown_runs(p, cell_val)
        else:
            p = doc.add_paragraph()
            _add_formatted_markdown_runs(p, trimmed)

        i += 1

    return doc


def save_markdown_to_docx_bytes(md_text: str) -> bytes:
    """Convert Markdown text to .docx and return raw bytes."""
    doc = create_docx_from_markdown(md_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def clean_extracted_pdf_markdown(md_text: str, extract_picture_text: bool = True) -> str:
    """
    Clean up artifact tags, picture text boxes, table cells, and HTML formatting from extracted PDF markdown.
    - Removes inline HTML tags like <mark>, <u>, <ins>, <span>, <font>, <del>, <strike> while preserving their text.
    - If extract_picture_text is False: Completely strips <!-- Start/End of picture text --> blocks.
    - If extract_picture_text is True: Normalizes picture text blocks, breaks <br>, and repairs glued PascalCase words ONLY inside picture blocks.
    - Formats markdown tables so rows remain single-line, bullet items are cleanly formatted, and trailing dashes are stripped.
    - Leaves normal body text (e.g. PowerPoint, LinkedIn, GmbH, ISO27001) 100% untouched.
    """
    if not md_text:
        return ""

    text = md_text

    # 1. Remove inline HTML markup tags (e.g. <mark>, <u>, <ins>, <span>, etc.) preserving the text inside
    text = strip_html_markup(text)

    # 2. Handle picture/graphics text blocks
    if not extract_picture_text:
        # Strip all picture text blocks entirely
        text = re.sub(r"<!--\s*Start of picture text\s*-->.*?<!--\s*End of picture text\s*-->", "", text, flags=re.DOTALL | re.IGNORECASE)
    else:
        def _clean_picture_block(match: re.Match) -> str:
            block = match.group(0)
            # Remove comment markers
            block = re.sub(r"<!--\s*(?:Start|End) of picture text\s*-->", "\n\n", block, flags=re.IGNORECASE)
            # Convert <br> tags in picture block to clean newlines
            block = re.sub(r"<br\s*/?>", "\n", block, flags=re.IGNORECASE)
            # Separate glued PascalCase words ONLY inside picture blocks (e.g. WandhovenWolfgang -> Wandhoven Wolfgang)
            def _separate_pascal_case(m: re.Match) -> str:
                s = m.group(0)
                if "http" in s or "www." in s or "/" in s:
                    return s
                return re.sub(r"(?<=[a-zäöüß])(?=[A-ZÄÖÜ])", " ", s)
            block = re.sub(r"\b[A-Za-zÄÖÜäöüß]{4,}\b", _separate_pascal_case, block)
            return block

        text = re.sub(r"<!--\s*Start of picture text\s*-->.*?<!--\s*End of picture text\s*-->", _clean_picture_block, text, flags=re.DOTALL | re.IGNORECASE)

    # 3. Clean markdown tables specifically so cell linebreaks (<br>) and list bullets are preserved without breaking table rows
    lines = text.splitlines()
    cleaned_lines = []
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 3:
            parts = stripped.split("|")
            cells = parts[1:-1]
            if all(re.match(r"^:?-+:?$", c.strip()) for c in cells if c.strip()):
                cleaned_lines.append("| " + " | ".join(c.strip() for c in cells) + " |")
                continue
            cleaned_cells = [clean_markdown_table_cell(c) for c in cells]

            # Heal split table cells from indented TOC (Inhaltsverzeichnis) columns
            if len(cleaned_cells) == 4:
                col0, col1, col2, col3 = cleaned_cells
                if (
                    re.match(r"^(?:\d+|[A-ZIVXLCDM]+)$", col0)
                    and re.match(r"^[A-Za-zÄÖÜäöüß]+$", col1)
                    and re.match(r"^[a-zäöüß]", col2)
                    and re.match(r"^(?:\d+|[ivxlcdmIVXLCDM]+|[-–—])$", col3)
                ):
                    cleaned_cells = [col0, "", col1 + col2, col3]

            cleaned_lines.append("| " + " | ".join(cleaned_cells) + " |")
        else:
            # Outside tables: convert any remaining <br> to newlines, but DO NOT run PascalCase separation on normal text!
            l = re.sub(r"<br\s*/?>", "\n", stripped, flags=re.IGNORECASE)
            cleaned_lines.append(l)

    text = "\n".join(cleaned_lines)

    # 4. Clean any remaining orphan comments or brackets
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)

    # 5. Normalize multiple whitespace and excessive newlines
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _extract_single_page_in_memory(
    raw_bytes: bytes,
    page_idx: int,
    margins: tuple,
) -> Tuple[int, str]:
    """
    Extract structured Markdown text for a single page directly from in-memory PDF bytes.
    Each worker thread opens its own PyMuPDF document instance and closes it in a finally block,
    ensuring 100% thread safety with zero disk I/O.
    """
    import pymupdf
    import pymupdf4llm.helpers.pymupdf_rag as rag

    doc = pymupdf.open(stream=raw_bytes, filetype="pdf")
    try:
        try:
            p_md = rag.to_markdown(
                doc,
                pages=[page_idx],
                margins=margins,
                table_strategy="lines",
            ).strip()
        except Exception:
            p_md = doc[page_idx].get_text("text").strip()
        return page_idx, p_md
    except Exception as ex:
        return page_idx, f"[Hinweis: Seite {page_idx + 1} konnte nicht extrahiert werden: {ex}]"
    finally:
        doc.close()


def extract_text_from_pdf_bytes(
    raw_bytes: bytes,
    filename: str = "document.pdf",
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
    include_headers_footers: bool = False,
    extract_picture_text: bool = True,
) -> str:
    """
    Extract structured Markdown text from PDF bytes using PyMuPDF's RAG layout engine.
    For multi-page PDFs, utilizes in-memory multi-threading (ThreadPoolExecutor) with zero disk I/O,
    eliminating temporary file race conditions, cross-process startup lag, and delivering smooth
    real-time page-by-page progress updates.
    Preserves headings, lists, tables, and bold/italic styles while maintaining intact word and sentence structure.
    Supports optional progress_callback(current_page, total_pages, status_text) for large PDFs.
    Preserves document title on page 1 while suppressing running headers/footers on subsequent pages when include_headers_footers=False.
    Raises ValueError if PDF contains pages but zero extractable text (e.g. scanned image PDF).
    """
    import pymupdf

    doc = pymupdf.open(stream=raw_bytes, filetype="pdf")
    doc_pages = doc.page_count

    # Check if there is any extractable text across pages
    has_text = any(page.get_text().strip() for page in doc)
    doc.close()

    if doc_pages > 0 and not has_text:
        raise ValueError(
            f"No extractable text found in PDF '{filename}' ({doc_pages} pages). "
            f"The file appears to be a scanned/image-based PDF requiring OCR, or an empty document."
        )

    if doc_pages == 0:
        return ""

    if doc_pages == 1:
        # Single page: Execute directly in memory
        if progress_callback:
            progress_callback(1, 1, "PDF-Inhalt wird extrahiert...")
        margins = (0, 0, 0, 0) if include_headers_footers else (0, 0, 0, 50)
        _, p_md = _extract_single_page_in_memory(raw_bytes, 0, margins)
        page_mds = [p_md]
    else:
        # Multi-page: Execute concurrently in memory using ThreadPoolExecutor
        cpu_cores = os.cpu_count() or 4
        max_workers = min(cpu_cores, 6, doc_pages)
        page_mds = [None] * doc_pages
        completed_pages = 0

        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            futs = []
            for page_idx in range(doc_pages):
                if include_headers_footers:
                    margins = (0, 0, 0, 0)
                else:
                    margins = (0, 0, 0, 50) if page_idx == 0 else (0, 50, 0, 50)

                futs.append(
                    executor.submit(
                        _extract_single_page_in_memory,
                        raw_bytes,
                        page_idx,
                        margins,
                    )
                )

            for fut in concurrent.futures.as_completed(futs):
                try:
                    p_idx, page_text = fut.result()
                    page_mds[p_idx] = page_text
                    completed_pages += 1
                except Exception as ex:
                    logging.error(f"PDF extraction thread error: {ex}")

                if progress_callback:
                    progress_callback(
                        completed_pages,
                        doc_pages,
                        f"PDF-Seite {completed_pages} von {doc_pages} extrahiert...",
                    )

    md_text = "\n\n".join(p for p in page_mds if p)
    return clean_extracted_pdf_markdown(md_text, extract_picture_text=extract_picture_text)


def extract_text_from_pdf(
    path: Union[str, Path],
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
    include_headers_footers: bool = False,
    extract_picture_text: bool = True,
) -> str:
    """
    Extract structured Markdown text from PDF files using pymupdf RAG layout engine.
    Raises ValueError if PDF contains pages but zero extractable text (e.g. scanned image PDF).
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")
    raw_data = safe_read_bytes(p)
    return extract_text_from_pdf_bytes(
        raw_data,
        filename=p.name,
        progress_callback=progress_callback,
        include_headers_footers=include_headers_footers,
        extract_picture_text=extract_picture_text,
    )


def read_document_from_bytes(
    data: bytes,
    filename: str,
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
    include_headers_footers: bool = False,
    extract_picture_text: bool = True,
) -> str:
    """Unified document reader from in-memory bytes with optional progress callback and extraction options."""
    ext = Path(filename).suffix.lower()
    if ext in [".txt", ".md"]:
        if progress_callback:
            progress_callback(1, 1, "Textdatei wird eingelesen...")
        text = extract_text_from_txt_bytes(data)
        return strip_html_markup(text)
    elif ext == ".json":
        if progress_callback:
            progress_callback(1, 1, "JSON-Struktur wird eingelesen...")
        return extract_text_from_json_bytes(data)
    elif ext == ".csv":
        if progress_callback:
            progress_callback(1, 1, "CSV-Tabelle wird eingelesen...")
        return extract_text_from_csv_bytes(data)
    elif ext == ".docx":
        if progress_callback:
            progress_callback(1, 1, "Word-Dokumentstruktur wird analysiert...")
        return extract_text_from_docx_bytes(data, include_headers_footers=include_headers_footers)
    elif ext == ".pdf":
        return extract_text_from_pdf_bytes(
            data,
            filename=filename,
            progress_callback=progress_callback,
            include_headers_footers=include_headers_footers,
            extract_picture_text=extract_picture_text,
        )
    else:
        raise UnsupportedFileFormatError(
            f"Unsupported file format: '{ext}'. Supported formats: .txt, .md, .json, .csv, .docx, .pdf"
        )


def read_document(
    file_path: Union[str, Path],
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
    include_headers_footers: bool = False,
    extract_picture_text: bool = True,
) -> str:
    """Unified document reader supporting .txt, .md, .json, .csv, .docx, and .pdf."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    raw_data = safe_read_bytes(path)
    return read_document_from_bytes(
        raw_data,
        path.name,
        progress_callback=progress_callback,
        include_headers_footers=include_headers_footers,
        extract_picture_text=extract_picture_text,
    )
