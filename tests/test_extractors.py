import pytest
from local_anonymizer.extractors import read_document, extract_text_from_pdf, extract_text_from_txt, UnsupportedFileFormatError
from pathlib import Path

def test_scanned_pdf_error_handling(tmp_path, mocker):
    # Mock pymupdf to return empty text
    mock_doc = mocker.MagicMock()
    mock_doc.__enter__.return_value = mock_doc
    mock_page = mocker.MagicMock()
    mock_page.get_text.return_value = "   \n"
    mock_doc.__iter__.return_value = [mock_page]
    mock_doc.page_count = 1
    
    mocker.patch("pymupdf.open", return_value=mock_doc)
    dummy_pdf = tmp_path / "dummy.pdf"
    dummy_pdf.write_bytes(b"%PDF-1.4 dummy")
    
    with pytest.raises(ValueError, match="scanned/image-based PDF requiring OCR"):
        extract_text_from_pdf(dummy_pdf)

def test_multi_encoding_txt(tmp_path):
    txt_file = tmp_path / "test.txt"
    # Write in cp1252
    txt_file.write_bytes("Müller äöü €".encode("cp1252"))
    
    text = extract_text_from_txt(txt_file)
    assert "Müller äöü €" in text

def test_unsupported_file_format(tmp_path):
    unknown_file = tmp_path / "dummy.unknown"
    unknown_file.write_text("content", encoding="utf-8")
    with pytest.raises(UnsupportedFileFormatError):
        read_document(unknown_file)


def test_pdf_to_markdown_structure(tmp_path):
    import pymupdf
    pdf_path = tmp_path / "sample_struct.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Hauptueberschrift", fontsize=24)
    page.insert_text((50, 100), "Dies ist ein normaler Textabsatz.", fontsize=11)
    doc.save(str(pdf_path))
    doc.close()

    extracted = extract_text_from_pdf(pdf_path)
    assert "Hauptueberschrift" in extracted
    assert "Dies ist ein normaler Textabsatz." in extracted


def test_anonymize_within_docx_markdown_headings_preserves_structure(tmp_path):
    import docx
    from local_anonymizer.extractors import extract_text_from_docx
    from local_anonymizer.anonymizer import LocalAnonymizer

    docx_path = tmp_path / "headings.docx"
    doc = docx.Document()
    doc.add_heading("Projektleiter Julia Meier Bericht", level=1)
    doc.save(str(docx_path))

    extracted = extract_text_from_docx(docx_path)
    assert "# Projektleiter Julia Meier Bericht" in extracted

    anonymizer = LocalAnonymizer()
    roles = {"julia meier": "STUDENT"}
    result = anonymizer.anonymize(extracted, format_mode="numbered_role", roles=roles)

    assert "# Projektleiter [PERSON_1_STUDENT] Bericht" in result.anonymized_text
    
    restored = LocalAnonymizer.de_anonymize(result.anonymized_text, result.mapping)
    assert restored == extracted


def test_anonymize_within_markdown_tables_preserves_structure():
    from local_anonymizer.anonymizer import LocalAnonymizer

    anonymizer = LocalAnonymizer()
    table_md = (
        "| Name | Rolle | Organisation |\n"
        "| :--- | :--- | :--- |\n"
        "| Julia Meier | Studentin | ZHAW |\n"
        "| Max Mustermann | Dozent | ETH |\n"
    )
    roles = {"julia meier": "STUDENT", "max mustermann": "DOZENT"}
    result = anonymizer.anonymize(table_md, format_mode="numbered_role", roles=roles)

    assert "[PERSON_1_STUDENT]" in result.anonymized_text
    assert "[PERSON_2_DOZENT]" in result.anonymized_text
    lines = result.anonymized_text.strip().splitlines()
    assert len(lines) == 4
    for line in lines:
        assert line.startswith("|") and line.endswith("|")

    # Roundtrip check
    restored = LocalAnonymizer.de_anonymize(result.anonymized_text, result.mapping)
    assert restored == table_md


def test_create_docx_from_markdown_with_native_styles():
    import docx
    from local_anonymizer.extractors import create_docx_from_markdown, save_markdown_to_docx_bytes

    md_content = (
        "# Hauptkapitel\n\n"
        "Dies ist ein normaler Textabsatz mit Informationen.\n\n"
        "## Unterabschnitt\n\n"
        "- Punkt 1\n"
        "- Punkt 2\n\n"
        "| Spalte A | Spalte B |\n"
        "| :--- | :--- |\n"
        "| Wert 1 | Wert 2 |\n"
    )

    doc = create_docx_from_markdown(md_content)
    # Verify paragraph styles and text
    paras = [p for p in doc.paragraphs if p.text]
    assert len(paras) >= 4
    assert paras[0].text == "Hauptkapitel"
    assert "Heading 1" in paras[0].style.name or "Überschrift 1" in paras[0].style.name or "heading 1" in paras[0].style.name.lower()
    assert paras[1].text == "Dies ist ein normaler Textabsatz mit Informationen."
    assert paras[2].text == "Unterabschnitt"
    assert "Heading 2" in paras[2].style.name or "Überschrift 2" in paras[2].style.name or "heading 2" in paras[2].style.name.lower()

    # Verify table
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Spalte A"
    assert doc.tables[0].cell(1, 0).text == "Wert 1"

    # Verify raw bytes generation
    raw_bytes = save_markdown_to_docx_bytes(md_content)
    assert len(raw_bytes) > 0


def test_create_docx_from_markdown_inline_formatting():
    from local_anonymizer.extractors import create_docx_from_markdown

    md = "Hier ist **fetter Text**, *kursiver Text*, ***fett-kursiv*** und `Consolas-Code`."
    doc = create_docx_from_markdown(md)
    p = doc.paragraphs[0]

    assert "**" not in p.text
    assert "*" not in p.text
    assert "`" not in p.text
    assert p.text == "Hier ist fetter Text, kursiver Text, fett-kursiv und Consolas-Code."

    bold_runs = [r for r in p.runs if r.bold and not r.italic]
    assert any("fetter Text" in r.text for r in bold_runs)

    italic_runs = [r for r in p.runs if r.italic and not r.bold]
    assert any("kursiver Text" in r.text for r in italic_runs)

    bi_runs = [r for r in p.runs if r.bold and r.italic]
    assert any("fett-kursiv" in r.text for r in bi_runs)

    code_runs = [r for r in p.runs if r.font.name == "Consolas"]
    assert any("Consolas-Code" in r.text for r in code_runs)


def test_docx_custom_outline_levels_and_lists(tmp_path):
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from local_anonymizer.extractors import extract_text_from_docx

    docx_path = tmp_path / "custom_doc.docx"
    doc = docx.Document()

    # 1. Custom heading via outline level
    p_h1 = doc.add_paragraph("Firmeninterne Hauptueberschrift")
    pPr = p_h1._p.get_or_add_pPr()
    outlineLvl = OxmlElement("w:outlineLvl")
    outlineLvl.set(qn("w:val"), "0")
    pPr.append(outlineLvl)

    # 2. German Heading 2 style
    p_h2 = doc.add_heading("Zweiter Unterabschnitt", level=2)

    # 3. Bullet item
    doc.add_paragraph("Erster Aufzählungspunkt", style="List Bullet")

    # 4. Numbered item
    doc.add_paragraph("Nummerierter Punkt Eins", style="List Number")

    doc.save(str(docx_path))

    extracted = extract_text_from_docx(docx_path)
    assert "# Firmeninterne Hauptueberschrift" in extracted
    assert "## Zweiter Unterabschnitt" in extracted
    assert "- Erster Aufzählungspunkt" in extracted
    assert "1. Nummerierter Punkt Eins" in extracted


def test_docx_bold_paragraph_false_positive_prevention(tmp_path):
    import docx
    from docx.shared import Pt
    from local_anonymizer.extractors import extract_text_from_docx

    docx_path = tmp_path / "warning_doc.docx"
    doc = docx.Document()

    # Regular paragraph with bold text that is NOT a heading (has full sentence with period, normal size)
    p_warn = doc.add_paragraph()
    run = p_warn.add_run("WICHTIGER HINWEIS: Dieser Text ist fett und wichtig, aber ein normaler Fließtextabsatz.")
    run.bold = True
    run.font.size = Pt(11)

    doc.save(str(docx_path))

    extracted = extract_text_from_docx(docx_path)
    # Must NOT start with '#'
    assert not extracted.startswith("#")
    assert "WICHTIGER HINWEIS: Dieser Text ist fett und wichtig, aber ein normaler Fließtextabsatz." in extracted


def test_anonymize_within_csv_markdown_tables_preserves_structure(tmp_path):
    from local_anonymizer.extractors import read_document
    from local_anonymizer.anonymizer import LocalAnonymizer

    csv_path = tmp_path / "team.csv"
    csv_path.write_text(
        "Name,Rolle,Hochschule\n"
        "Julia Meier,Studentin,ZHAW\n"
        "Max Mustermann,Dozent,ETH\n",
        encoding="utf-8"
    )

    extracted_md = read_document(csv_path)
    # Verify CSV was parsed into markdown table
    assert "| Name | Rolle | Hochschule |" in extracted_md
    assert "| :--- | :--- | :--- |" in extracted_md
    assert "| Julia Meier | Studentin | ZHAW |" in extracted_md

    anonymizer = LocalAnonymizer()
    roles = {"julia meier": "STUDENT", "max mustermann": "DOZENT"}
    result = anonymizer.anonymize(extracted_md, format_mode="numbered_role", roles=roles)

    assert "[PERSON_1_STUDENT]" in result.anonymized_text
    assert "[PERSON_2_DOZENT]" in result.anonymized_text
    assert "[ORGANIZATION_1]" in result.anonymized_text or "ZHAW" in result.anonymized_text

    # Structure check: every line should be a markdown table row
    for line in result.anonymized_text.strip().splitlines():
        assert line.startswith("|") and line.endswith("|")

    # Roundtrip reversibility check
    restored = LocalAnonymizer.de_anonymize(result.anonymized_text, result.mapping)
    assert restored == extracted_md


def test_read_document_from_bytes_with_progress():
    import pymupdf
    from local_anonymizer.extractors import read_document_from_bytes

    doc = pymupdf.open()
    p1 = doc.new_page()
    p1.insert_text((50, 150), "Page 1 Content")
    p2 = doc.new_page()
    p2.insert_text((50, 150), "Page 2 Content")
    pdf_bytes = doc.tobytes()
    doc.close()

    progress_calls = []
    def on_progress(curr, total, msg):
        progress_calls.append((curr, total, msg))

    text = read_document_from_bytes(pdf_bytes, "test.pdf", progress_callback=on_progress)
    assert "Page 1 Content" in text
    assert "Page 2 Content" in text
    assert len(progress_calls) >= 2
    assert progress_calls[-1][0] == 2
    assert progress_calls[-1][1] == 2


def test_clean_extracted_pdf_markdown():
    from local_anonymizer.extractors import clean_extracted_pdf_markdown

    sample = (
        "## Dokumenttitel\n\n"
        "<mark>Dr. Andreas Schönenberger</mark> leitet das Projekt.\n"
        "Quelle: <u>www.sanitas.com</u>\n\n"
        "<!-- Start of picture text -->\n"
        "CEO<br>Dr. Andreas<br>Schönenberger<br>CMO / New<br>Corporate  CFO<br>\n"
        "Elias Frühauf Jan Schultz Kaspar  WandhovenWolfgang  Tobias Caluori<br>Trachsel<br>\n"
        "<!-- End of picture text -->"
    )

    # 1. With extract_picture_text=True
    cleaned = clean_extracted_pdf_markdown(sample, extract_picture_text=True)
    assert "<mark>" not in cleaned and "</mark>" not in cleaned
    assert "<u>" not in cleaned and "</u>" not in cleaned
    assert "Dr. Andreas Schönenberger leitet das Projekt." in cleaned
    assert "Quelle: www.sanitas.com" in cleaned
    assert "<!-- Start of picture text -->" not in cleaned
    assert "<br>" not in cleaned
    assert "Wandhoven Wolfgang" in cleaned

    # 2. With extract_picture_text=False (picture text completely omitted)
    cleaned_no_pic = clean_extracted_pdf_markdown(sample, extract_picture_text=False)
    assert "Dr. Andreas Schönenberger leitet das Projekt." in cleaned_no_pic
    assert "Quelle: www.sanitas.com" in cleaned_no_pic
    assert "Wandhoven" not in cleaned_no_pic
    assert "Jan Schultz" not in cleaned_no_pic


def test_clean_extracted_pdf_markdown_sipoc_table():
    from local_anonymizer.extractors import clean_extracted_pdf_markdown

    raw_table = (
        "|**Supplier **|**Input**|**Process**|**Output**|**Customer **|\n"
        "|---|---|---|---|---|\n"
        "|-<br>Veranlasser:in<br>-<br>Leistungserbringer:in<br>-<br>SPS (Dienstleister von<br>Sanitas für Dokumenten-<br>scanning)|"
        "-<br>Rezept für ein Medikament<br>-<br>DMS-Ablage, Tarif-<br>und Deckungsregeln<br>-<br>-<br>-|"
        "Rezept wird von<br>Versicherungsnehmer:in bei<br>Veranlasser:in abgeholt<br><br>Dokumente werden bei<br>Eingang gescannt|"
        "-<br>Geprüfter Leistungs-<br>fall zu einem Rezept<br>-|"
        "Versicherungsnehmer:in|"
    )

    cleaned = clean_extracted_pdf_markdown(raw_table)

    # Validate that table rows remain single-line (exactly 3 lines: header, separator, row)
    assert len(cleaned.splitlines()) == 3
    # Validate broken hyphenation is fixed
    assert "Dokumentenscanning" in cleaned
    assert "Tarif- und Deckungsregeln" in cleaned
    assert "Leistungsfall" in cleaned
    # Validate list items are formatted with bullets without broken linebreaks
    assert "- Veranlasser:in" in cleaned
    assert "- Leistungserbringer:in" in cleaned
    # Validate trailing empty dashes are removed
    assert "<br>-<br>-" not in cleaned
    assert "zu einem Rezept |" in cleaned or "zu einem Rezept<br>" in cleaned or "zu einem Rezept" in cleaned


def test_pdf_headers_footers_toggle():
    import pymupdf
    from local_anonymizer.extractors import extract_text_from_pdf_bytes

    doc = pymupdf.open()
    # Page 1: Has document title at top and body and footer
    p1 = doc.new_page(width=595, height=842)
    p1.insert_text((50, 30), "Dokumententitel auf Seite 1", fontsize=16)
    p1.insert_text((50, 150), "Haupttext des Dokuments auf Seite 1.", fontsize=11)
    p1.insert_text((50, 800), "Seite 1 von 2 - Fußzeile", fontsize=9)

    # Page 2: Has running header at top, body and footer
    p2 = doc.new_page(width=595, height=842)
    p2.insert_text((50, 30), "Wiederkehrende Kopfzeile Sanitas", fontsize=9)
    p2.insert_text((50, 150), "Haupttext des Dokuments auf Seite 2.", fontsize=11)
    p2.insert_text((50, 800), "Seite 2 von 2 - Fußzeile", fontsize=9)

    pdf_bytes = doc.tobytes()
    doc.close()

    # Default (include_headers_footers=False):
    # - Page 1 Title protected (not lost!)
    # - Running header on Page 2 filtered out
    # - Footers on all pages filtered out
    text_no_hf = extract_text_from_pdf_bytes(pdf_bytes, include_headers_footers=False)
    assert "Dokumententitel auf Seite 1" in text_no_hf
    assert "Haupttext des Dokuments auf Seite 1." in text_no_hf
    assert "Haupttext des Dokuments auf Seite 2." in text_no_hf
    assert "Wiederkehrende Kopfzeile Sanitas" not in text_no_hf
    assert "Fußzeile" not in text_no_hf

    # With headers and footers (include_headers_footers=True): Everything included
    text_with_hf = extract_text_from_pdf_bytes(pdf_bytes, include_headers_footers=True)
    assert "Dokumententitel auf Seite 1" in text_with_hf
    assert "Haupttext des Dokuments auf Seite 1." in text_with_hf
    assert "Wiederkehrende Kopfzeile Sanitas" in text_with_hf
    assert "Fußzeile" in text_with_hf


def test_docx_headers_footers_toggle():
    import io
    import docx
    from local_anonymizer.extractors import extract_text_from_docx_bytes

    doc = docx.Document()
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "Header Text in Word"
    doc.add_paragraph("Body paragraph in Word.")
    sec.footer.paragraphs[0].text = "Footer Text in Word"

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    # Default (include_headers_footers=False)
    text_no_hf = extract_text_from_docx_bytes(docx_bytes, include_headers_footers=False)
    assert "Body paragraph in Word." in text_no_hf
    assert "Header Text in Word" not in text_no_hf
    assert "Footer Text in Word" not in text_no_hf

    # With headers and footers (include_headers_footers=True)
    text_with_hf = extract_text_from_docx_bytes(docx_bytes, include_headers_footers=True)
    assert "Body paragraph in Word." in text_with_hf
    assert "Header Text in Word" in text_with_hf
    assert "Footer Text in Word" in text_with_hf


def test_clean_extracted_pdf_markdown_pascal_case_flow_text_preserved():
    """Verify PascalCase regex ONLY splits picture text blocks and does NOT alter standard flow text."""
    from local_anonymizer.extractors import clean_extracted_pdf_markdown

    sample = (
        "Hier ist ein normaler Fließtext mit PowerPoint, LinkedIn, GmbH, iPhone und Sanitas.\n\n"
        "<!-- Start of picture text -->\n"
        "Dr. Andreas Schönenberger CMO / New Corporate CFO Operation , IT Center Business Offering\n"
        "Jan Schultz Kaspar WandhovenWolfgang Tobias Caluori Trachsel\n"
        "<!-- End of picture text -->\n\n"
        "Abschlussbericht von Sanitas GmbH mit PowerPoint Präsentation."
    )

    cleaned = clean_extracted_pdf_markdown(sample, extract_picture_text=True)

    # Standard flow text must NOT be altered
    assert "PowerPoint" in cleaned
    assert "Power Point" not in cleaned
    assert "LinkedIn" in cleaned
    assert "Linked In" not in cleaned
    assert "GmbH" in cleaned
    assert "Gmb H" not in cleaned
    assert "iPhone" in cleaned

    # Picture text block MUST be properly split
    assert "Wandhoven Wolfgang" in cleaned
    assert "<!-- Start of picture text -->" not in cleaned


def test_extract_text_from_csv_bytes_multiline():
    """Verify quoted multiline CSV cells are correctly formatted into clean Markdown table cells."""
    from local_anonymizer.extractors import extract_text_from_csv_bytes

    csv_data = (
        'Name,"Rolle & Beschreibung",Status\n'
        'Julia Meier,"Projektleiterin\nZuständig für Release 1.0",Aktiv\n'
        'Max Muster,"Dozent\nFachbereich Informatik",Inaktiv\n'
    ).encode("utf-8")

    extracted = extract_text_from_csv_bytes(csv_data)

    # Must be valid Markdown table (Header, Separator, 2 Data Rows = 4 lines)
    lines = [l for l in extracted.strip().splitlines() if l.strip()]
    assert len(lines) == 4
    assert "| Name | Rolle & Beschreibung | Status |" in lines[0]
    assert "Projektleiterin<br>Zuständig für Release 1.0" in lines[2]
    assert "Dozent<br>Fachbereich Informatik" in lines[3]


def test_strip_html_markup():
    """Verify strip_html_markup cleans annotations while preserving content."""
    from local_anonymizer.extractors import strip_html_markup

    text = "Hier ist <mark>wichtiger Text</mark> und <u>unterstrichenes Wort</u> sowie <span style='color:red;'>roter Text</span>."
    cleaned = strip_html_markup(text)
    assert cleaned == "Hier ist wichtiger Text und unterstrichenes Wort sowie roter Text."


def test_clean_extracted_pdf_markdown_toc_table_healing():
    """Verify indented TOC table column splits (e.g. 'Proz | essidentifikation') are healed cleanly."""
    from local_anonymizer.extractors import clean_extracted_pdf_markdown

    raw_toc = (
        "# **Inhaltsverzeichnis**\n\n"
        "| 1 | Proz | essidentifikation | 1 |\n"
        "| --- | --- | --- | --- |\n"
        "| | 1.1 | Unternehmenskontext | 1 |\n"
        "| | 1.2 | Einbettung in Prozessarchitektur | 3 |\n"
        "| 2 | Proz | esserhebung | 6 |\n"
        "| | 2.1 | Verwendete Systeme | 6 |\n"
        "| 6 | Anh | ang | 14 |\n"
        "| | 6.1 | Glossar | 14 |\n"
    )

    cleaned = clean_extracted_pdf_markdown(raw_toc)

    assert "| 1 | | Prozessidentifikation | 1 |" in cleaned
    assert "| | 1.1 | Unternehmenskontext | 1 |" in cleaned
    assert "| 2 | | Prozesserhebung | 6 |" in cleaned
    assert "| 6 | | Anhang | 14 |" in cleaned
    assert "Proz | ess" not in cleaned
    assert "Anh | ang" not in cleaned


def test_pdf_parallel_extraction_ordering():
    """Verify that multi-page PDF multiprocessing extraction returns pages in exact numerical order and handles margins."""
    import pymupdf
    from local_anonymizer.extractors import extract_text_from_pdf_bytes

    doc = pymupdf.open()
    for i in range(1, 11):
        page = doc.new_page(width=595, height=842)
        if i == 1:
            page.insert_text((50, 30), "Document Title Page 1", fontsize=16)
        else:
            page.insert_text((50, 30), f"Header Page {i}", fontsize=9)
        page.insert_text((50, 200), f"Unique Content Section Page {i}", fontsize=12)
        page.insert_text((50, 810), f"Footer Page {i}", fontsize=9)

    pdf_bytes = doc.tobytes()
    doc.close()

    progress_steps = []
    def on_progress(curr, total, msg):
        progress_steps.append((curr, total))

    extracted = extract_text_from_pdf_bytes(pdf_bytes, progress_callback=on_progress)

    # 1. Verify exact page ordering
    last_pos = -1
    for i in range(1, 11):
        needle = f"Unique Content Section Page {i}"
        pos = extracted.find(needle)
        assert pos != -1, f"Page {i} content missing!"
        assert pos > last_pos, f"Page {i} appeared out of order!"
        last_pos = pos

    # 2. Verify Page 1 title preserved, while Page 2-10 running headers and footers are filtered
    assert "Document Title Page 1" in extracted
    assert "Header Page 2" not in extracted
    assert "Header Page 10" not in extracted
    assert "Footer Page 1" not in extracted
    assert "Footer Page 5" not in extracted

    # 3. Verify progress callback was triggered up to (10, 10)
    assert len(progress_steps) > 0
    assert progress_steps[-1] == (10, 10)


def test_pdf_parallel_extraction_corrupted_page_resilience(monkeypatch):
    """Verify that a failure on a single page falls back gracefully without aborting the entire document."""
    import pymupdf
    import pymupdf4llm.helpers.pymupdf_rag as rag
    from local_anonymizer.extractors import extract_text_from_pdf_bytes

    doc = pymupdf.open()
    for i in range(1, 5):
        page = doc.new_page(width=595, height=842)
        page.insert_text((50, 200), f"Resilient Section {i}", fontsize=12)

    pdf_bytes = doc.tobytes()
    doc.close()

    original_to_markdown = rag.to_markdown

    def faulty_to_markdown(doc, pages=None, **kwargs):
        if pages and pages[0] == 1:  # Page 2 (0-indexed 1) fails in markdown parser
            raise RuntimeError("Simulated broken layout on page 2")
        return original_to_markdown(doc, pages=pages, **kwargs)

    monkeypatch.setattr(rag, "to_markdown", faulty_to_markdown)

    extracted = extract_text_from_pdf_bytes(pdf_bytes)

    # Page 1, 3, 4 extracted normally
    assert "Resilient Section 1" in extracted
    assert "Resilient Section 3" in extracted
    assert "Resilient Section 4" in extracted
    # Page 2 still has text via stage 2 plain-text fallback!
    assert "Resilient Section 2" in extracted


def test_pdf_parallel_extraction_in_memory_no_temp_files():
    """Verify that multi-page PDF extraction runs 100% in-memory with ThreadPoolExecutor without creating temp files."""
    import pymupdf
    from local_anonymizer.extractors import extract_text_from_pdf_bytes, TEMP_UPLOADS_DIR

    doc = pymupdf.open()
    for i in range(1, 6):
        page = doc.new_page(width=595, height=842)
        page.insert_text((50, 200), f"In-Memory Section Page {i}", fontsize=12)
    pdf_bytes = doc.tobytes()
    doc.close()

    # Pre-check: count existing PDF files in TEMP_UPLOADS_DIR
    TEMP_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    initial_pdfs = list(TEMP_UPLOADS_DIR.glob("*.pdf"))

    text = extract_text_from_pdf_bytes(pdf_bytes)
    assert "In-Memory Section Page 1" in text
    assert "In-Memory Section Page 5" in text

    # Post-check: zero temporary files were created on disk
    remaining_pdfs = list(TEMP_UPLOADS_DIR.glob("*.pdf"))
    assert len(remaining_pdfs) == len(initial_pdfs)


def test_pdf_extraction_concurrent_normal_app_import_safety():
    """
    Regression test: Concurrent normal app import (or second app start without any worker env var)
    MUST NOT interfere with or delete active extraction files or in-memory jobs.
    """
    import subprocess
    import sys
    import os
    import pymupdf
    from pathlib import Path
    from local_anonymizer.extractors import extract_text_from_pdf_bytes

    doc = pymupdf.open()
    for i in range(1, 8):
        page = doc.new_page(width=595, height=842)
        page.insert_text((50, 200), f"Concurrent Robust Page {i}", fontsize=12)
    pdf_bytes = doc.tobytes()
    doc.close()

    # Simulate concurrent normal app import in another process
    test_script = (
        "import os, sys\n"
        "sys.path.insert(0, 'src')\n"
        "import local_anonymizer.extractors\n"
        "import app\n"
        "print('APP_IMPORT_CONCURRENT_SUCCESS')\n"
    )

    env = dict(os.environ)
    env.pop("LOCAL_ANONYMIZER_PDF_WORKER", None)
    env["PYTHONPATH"] = "src"

    res = subprocess.run(
        [sys.executable, "-c", test_script],
        env=env,
        capture_output=True,
        text=True,
        cwd=str(Path(__file__).parent.parent),
    )
    assert res.returncode == 0, f"App import failed: {res.stderr}"
    assert "APP_IMPORT_CONCURRENT_SUCCESS" in res.stdout

    # In-memory extraction finishes cleanly
    extracted = extract_text_from_pdf_bytes(pdf_bytes)
    assert "Concurrent Robust Page 1" in extracted
    assert "Concurrent Robust Page 7" in extracted


def test_age_based_temp_cleanup_preserves_recent_files():
    """Verify that age-based cleanup preserves recently created temp files and deletes only old ones."""
    import os
    import time
    import tempfile
    from local_anonymizer.extractors import cleanup_extraction_temp_files, TEMP_UPLOADS_DIR
    from app import cleanup_temp_uploads, UPLOAD_DIR

    TEMP_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(suffix=".pdf", dir=TEMP_UPLOADS_DIR, delete=False) as tmp_new:
        tmp_new.write(b"new active temp file")
        new_path = Path(tmp_new.name)

    with tempfile.NamedTemporaryFile(suffix=".pdf", dir=TEMP_UPLOADS_DIR, delete=False) as tmp_old:
        tmp_old.write(b"old stale temp file")
        old_path = Path(tmp_old.name)

    try:
        # Age old_path artificially by 3600s
        old_time = time.time() - 3600
        os.utime(old_path, (old_time, old_time))

        cleanup_extraction_temp_files(max_age_seconds=1800)

        assert new_path.exists(), "New temp file was unexpectedly deleted!"
        assert not old_path.exists(), "Old stale temp file was not cleaned up!"
    finally:
        new_path.unlink(missing_ok=True)
        old_path.unlink(missing_ok=True)

